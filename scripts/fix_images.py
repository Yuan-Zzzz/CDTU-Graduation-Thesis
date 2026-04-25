import sys
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_size, east_asia='宋体', ascii_font='Times New Roman'):
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:cs'), ascii_font)
    
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(font_size))
    
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(font_size))

def fix_paragraph_format(paragraph, font_size, alignment='center', line_spacing=360, space_after=None):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    
    pJc = pPr.find(qn('w:jc'))
    if pJc is None:
        pJc = OxmlElement('w:jc')
        pPr.append(pJc)
    pJc.set(qn('w:val'), alignment)
    
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(line_spacing))
    spacing.set(qn('w:lineRule'), 'auto')
    
    if space_after is not None:
        spacing.set(qn('w:after'), str(space_after))
    
    for run in paragraph.runs:
        set_run_font(run, font_size)

def is_figure_caption(text):
    return bool(re.match(r'^图\s*\d+', text.strip()))

def fix_images(filename):
    doc = Document(filename)
    body = doc._element.body
    children = list(body)
    
    figure_count = 0
    
    for i, child in enumerate(children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'p':
            text = ''
            for t in child.iter():
                if t.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                    text += t.text or ''
            text = text.strip()
            
            if is_figure_caption(text):
                figure_count += 1
                for para in doc.paragraphs:
                    if para._element is child:
                        fix_paragraph_format(para, 21, alignment='center', line_spacing=360)
                        
                        if re.match(r'^图\d+', text):
                            for run in para.runs:
                                run_text = run.text
                                if re.match(r'^图\d+', run_text):
                                    new_text = re.sub(r'(图\d+)', r'\1 ', run_text, count=1)
                                    run.text = new_text
                                    break
                        
                        print(f"修复图片标题: {text[:50]}")
                        break
            
            has_image = False
            for run in child.iter():
                if run.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                    drawings = run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    if drawings:
                        has_image = True
                        break
            
            if has_image:
                for para in doc.paragraphs:
                    if para._element is child:
                        fix_paragraph_format(para, 21, alignment='center', line_spacing=360)
                        
                        pPr = para._element.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            para._element.insert(0, pPr)
                        
                        keepNext = pPr.find(qn('w:keepNext'))
                        if keepNext is None:
                            keepNext = OxmlElement('w:keepNext')
                            pPr.append(keepNext)
                        
                        for drawing in para._element.iter('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            anchors = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                            for anchor in anchors:
                                inline = OxmlElement('wp:inline')
                                extent = anchor.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                docPr = anchor.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')
                                if extent is not None:
                                    inline.append(extent)
                                if docPr is not None:
                                    inline.append(docPr)
                                graphic = anchor.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
                                if graphic is not None:
                                    inline.append(graphic)
                                drawing.remove(anchor)
                                drawing.append(inline)
                        
                        print(f"修复图片段落: {text[:50] if text else '(empty)'}")
                        break
    
    doc.save(filename)
    print(f"图片格式修复完成: {filename} (共处理 {figure_count} 个图片标题)")

if __name__ == '__main__':
    fix_images(sys.argv[1])

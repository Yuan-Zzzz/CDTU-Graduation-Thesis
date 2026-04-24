from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.hdrftr import FooterPart, HeaderPart
import copy
import re
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
FINAL_DOC = REPO_ROOT / "output" / "Final_thesis.docx"

doc = Document(str(FINAL_DOC))
body = doc._element.body

def insert_page_break_before(element):
    prev = element.getprevious()
    while prev is not None:
        prev_tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
        if prev_tag == 'p':
            has_page_break = False
            for br in prev.iter():
                if br.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br':
                    if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                        has_page_break = True
                        break
            
            if has_page_break:
                return
            
            has_section_break = False
            pPr = prev.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                sectPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if sectPr is not None:
                    has_section_break = True
            
            if not has_section_break:
                break
        
        prev = prev.getprevious()
    
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    element.addprevious(pb_p)

sdt_element = None
for child in body:
    if child.tag.endswith('}sdt'):
        sdt_element = child
        break

if sdt_element is not None:
    insert_page_break_before(sdt_element)

for p in doc.paragraphs:
    text = p.text.strip()
    style = p.style.name if p.style else ''
    
    if text == '摘要' and '中文摘要' in style:
        insert_page_break_before(p._element)
    elif text == 'Abstract' and '英文摘要' in style:
        insert_page_break_before(p._element)
    elif text == '参考文献':
        insert_page_break_before(p._element)
    elif text == '致    谢' or text == '致谢':
        insert_page_break_before(p._element)

chap1_p = None
for p in doc.paragraphs:
    if p.text.startswith('第一章') or p.text.startswith('第 1 章'):
        chap1_p = p._element
        break

if chap1_p is not None:
    prev = chap1_p.getprevious()
    has_section_break = False
    while prev is not None:
        prev_tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
        if prev_tag == 'p':
            pPr = prev.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                sectPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if sectPr is not None:
                    has_section_break = True
                    break
        prev = prev.getprevious()
    
    if not has_section_break:
        sect_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        sect_p.append(pPr)
        
        doc_sectPr = body.sectPr
        new_sectPr = copy.deepcopy(doc_sectPr)
        
        type_el = new_sectPr.find(qn('w:type'))
        if type_el is None:
            type_el = OxmlElement('w:type')
            new_sectPr.append(type_el)
        type_el.set(qn('w:val'), 'nextPage')
        
        pgNumType_front = new_sectPr.find(qn('w:pgNumType'))
        if pgNumType_front is None:
            pgNumType_front = OxmlElement('w:pgNumType')
            new_sectPr.append(pgNumType_front)
        pgNumType_front.set(qn('w:fmt'), 'upperRoman')
        pgNumType_front.set(qn('w:start'), '1')
        
        pPr.append(new_sectPr)
        chap1_p.addprevious(sect_p)
        
        pgNumType_main = doc_sectPr.find(qn('w:pgNumType'))
        if pgNumType_main is None:
            pgNumType_main = OxmlElement('w:pgNumType')
            doc_sectPr.append(pgNumType_main)
        pgNumType_main.set(qn('w:fmt'), 'decimal')
        pgNumType_main.set(qn('w:start'), '1')

header_xml_str = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wpsCustomData="http://www.wps.cn/officeDocument/2013/wpsCustomData" mc:Ignorable="w14 w15 wp14">
  <w:p w14:paraId="46D6B860">
    <w:pPr>
      <w:pStyle w:val="77"/>
      <w:rPr>
        <w:color w:val="FF0000"/>
      </w:rPr>
    </w:pPr>
    <w:r>
      <w:t>成都工业学院本科毕业论文（设计）</w:t>
    </w:r>
  </w:p>
</w:hdr>'''

header_element = parse_xml(header_xml_str.encode('utf-8'))
new_partname = doc.part.package.next_partname(PackURI('/word/header%d.xml'))
new_header_part = HeaderPart(new_partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml', header_element, doc.part.package)
header_rel_id = doc.part.relate_to(new_header_part, RT.HEADER)

roman_footer_xml_str = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wpsCustomData="http://www.wps.cn/officeDocument/2013/wpsCustomData" mc:Ignorable="w14 w15 wp14">
  <w:p w14:paraId="499F95DC">
    <w:pPr>
      <w:pStyle w:val="12"/>
      <w:framePr w:wrap="auto" w:vAnchor="text" w:hAnchor="margin" w:xAlign="center" w:y="1"/>
      <w:rPr>
        <w:rStyle w:val="26"/>
      </w:rPr>
    </w:pPr>
    <w:r>
      <w:rPr>
        <w:rStyle w:val="26"/>
      </w:rPr>
      <w:fldChar w:fldCharType="begin"/>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rStyle w:val="26"/>
      </w:rPr>
      <w:instrText xml:space="preserve"> PAGE </w:instrText>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rStyle w:val="26"/>
      </w:rPr>
      <w:fldChar w:fldCharType="separate"/>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rStyle w:val="26"/>
      </w:rPr>
      <w:t>I</w:t>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rStyle w:val="26"/>
      </w:rPr>
      <w:fldChar w:fldCharType="end"/>
    </w:r>
  </w:p>
</w:ftr>'''

chinese_footer_xml_str = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:office" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wpsCustomData="http://www.wps.cn/officeDocument/2013/wpsCustomData" mc:Ignorable="w14 w15 wp14">
  <w:p w14:paraId="1B0ACECC">
    <w:pPr>
      <w:pStyle w:val="85"/>
      <w:jc w:val="center"/>
      <w:rPr>
        <w:color w:val="FF0000"/>
      </w:rPr>
    </w:pPr>
    <w:r>
      <w:t xml:space="preserve">第 </w:t>
    </w:r>
    <w:r>
      <w:fldChar w:fldCharType="begin"/>
    </w:r>
    <w:r>
      <w:instrText xml:space="preserve">PAGE   * MERGEFORMAT</w:instrText>
    </w:r>
    <w:r>
      <w:fldChar w:fldCharType="separate"/>
    </w:r>
    <w:r>
      <w:t>1</w:t>
    </w:r>
    <w:r>
      <w:fldChar w:fldCharType="end"/>
    </w:r>
    <w:r>
      <w:t xml:space="preserve"> 页</w:t>
    </w:r>
  </w:p>
</w:ftr>'''

roman_footer_element = parse_xml(roman_footer_xml_str.encode('utf-8'))
roman_partname = doc.part.package.next_partname(PackURI('/word/footer%d.xml'))
roman_footer_part = FooterPart(roman_partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml', roman_footer_element, doc.part.package)
roman_footer_rel_id = doc.part.relate_to(roman_footer_part, RT.FOOTER)

chinese_footer_element = parse_xml(chinese_footer_xml_str.encode('utf-8'))
chinese_partname = doc.part.package.next_partname(PackURI('/word/footer%d.xml'))
chinese_footer_part = FooterPart(chinese_partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml', chinese_footer_element, doc.part.package)
chinese_footer_rel_id = doc.part.relate_to(chinese_footer_part, RT.FOOTER)

for i in range(len(doc.sections)):
    sectPr = doc.sections[i]._sectPr
    
    if i == 0:
        footer_rel_id = roman_footer_rel_id
    else:
        footer_rel_id = chinese_footer_rel_id
        header_ref = sectPr.find(qn('w:headerReference'))
        if header_ref is None:
            header_ref = OxmlElement('w:headerReference')
            sectPr.insert(0, header_ref)
        header_ref.set(qn('w:type'), 'default')
        header_ref.set(qn('r:id'), header_rel_id)
    
    footer_refs = sectPr.findall(qn('w:footerReference'))
    for ref in footer_refs:
        if ref.get(qn('w:type')) == 'default':
            ref.set(qn('r:id'), footer_rel_id)

headings = []
for para in doc.paragraphs:
    text = para.text.strip()
    is_chapter = False
    level = 1
    
    if text.startswith('第') and '章' in text:
        is_chapter = True
        level = 1
    elif re.match(r'^\d+\.\d+\.\d+', text):
        is_chapter = True
        level = 3
    elif re.match(r'^\d+\.\d+', text):
        is_chapter = True
        level = 2
    
    if is_chapter:
        headings.append({'text': text, 'level': level})

heading_page_map = {}
current_page = 1
lines_on_page = 0
MAX_LINES_PER_PAGE = 45

for para in doc.paragraphs:
    p = para._element
    text = para.text.strip()
    
    for br in p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
        if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
            current_page += 1
            lines_on_page = 0
    
    pPr = p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        sectPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
        if sectPr is not None:
            current_page += 1
            lines_on_page = 0
    
    if text:
        char_count = len(text)
        estimated_lines = max(1, (char_count + 39) // 40)
        lines_on_page += estimated_lines
        
        if lines_on_page > MAX_LINES_PER_PAGE:
            current_page += 1
            lines_on_page = estimated_lines
    
    for heading in headings:
        if heading['text'] == text:
            heading_page_map[heading['text']] = current_page
            break

for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'sdt':
        for sdt_child in child:
            sdt_tag = sdt_child.tag.split('}')[-1] if '}' in sdt_child.tag else sdt_child.tag
            if sdt_tag == 'sdtContent':
                for content_child in sdt_child:
                    content_tag = content_child.tag.split('}')[-1] if '}' in content_child.tag else content_child.tag
                    if content_tag == 'p':
                        text = ''
                        for t in content_child.iter():
                            if t.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                                text += t.text or ''
                        
                        heading_text = text.rstrip('0123456789').strip()
                        for heading in headings:
                            if heading['text'] == heading_text:
                                page_num = heading_page_map.get(heading['text'], 1)
                                runs = list(content_child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'))
                                if runs:
                                    last_run = runs[-1]
                                    t = last_run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                    if t is not None:
                                        t.text = str(page_num)
                                break
        break

doc.save(str(FINAL_DOC))

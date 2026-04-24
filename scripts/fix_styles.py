import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

filename = sys.argv[1]
doc = Document(filename)

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
TEMPLATE_PATH = REPO_ROOT / "assets" / "template.docx"

STYLE_MAPPING = {
    'Heading 1': '73',
    'Heading 2': '74',
    'Heading 3': '75',
    'FirstParagraph': '76',
    'BodyText': '76',
    'Compact': '76',
    'Bibliography': '79',
}

TOC_STYLE_MAPPING = {
    'TOCHeading': '67',
    'toc 1': '68',
    'toc 2': '69',
    'toc 3': '70',
}

TOC_STYLE_ID_MAPPING = {
    '14': '68',
    '17': '69',
    '8': '70',
}

TOC_LEVEL_STYLE_MAPPING = {
    'Heading 1': '68',
    'Heading 2': '69',
    'Heading 3': '70',
}

ABSTRACT_TITLE_STYLES = {
    '摘要': '61',
    'Abstract': '64',
}

ABSTRACT_CONTENT_STYLES = {
    '中文摘要 内容': '62',
    '英文摘要 内容': '65',
}

ABSTRACT_KEYWORD_STYLES = {
    '关键词': '63',
    'Keywords': '66',
}

def get_style_id_by_name(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style.style_id
    return None

def get_template_toc_style_ids(template_path):
    import zipfile
    import xml.etree.ElementTree as ET
    
    toc_styles = {}
    try:
        with zipfile.ZipFile(template_path, 'r') as z:
            with z.open('word/styles.xml') as f:
                styles_content = f.read().decode('utf-8')
        
        root = ET.fromstring(styles_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        for style in root.findall('.//w:style', ns):
            style_id = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
            name = style.find('w:name', ns)
            if name is not None:
                name_val = name.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if name_val in ['目录 一级标题', '目录 二级标题', '目录 三级标题']:
                    toc_styles[name_val] = style_id
    except Exception as e:
        print(f"读取模板样式失败: {e}")
    
    return toc_styles

def apply_style_by_id(paragraph, style_id):
    pPr = paragraph._element.get_or_add_pPr()
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        if len(pPr) > 0:
            pPr.insert(0, pStyle)
        else:
            pPr.append(pStyle)
    pStyle.set(qn('w:val'), style_id)

def is_chinese_text(text):
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            return True
    return False

def detect_abstract_context(paragraphs, current_idx):
    if current_idx < 0 or current_idx >= len(paragraphs):
        return None
    
    p = paragraphs[current_idx]
    text = p.text.strip()
    
    if text in ABSTRACT_TITLE_STYLES:
        return 'title', ABSTRACT_TITLE_STYLES[text]
    
    if text.startswith('关键词：') or text.startswith('Keywords:'):
        return 'keywords', ABSTRACT_KEYWORD_STYLES.get('关键词' if text.startswith('关键词') else 'Keywords')

    return None


def split_keyword_paragraph(paragraph):
    text = paragraph.text
    if text.startswith('关键词：'):
        label = '关键词：'
        content = text[len(label):]
    elif text.startswith('Keywords:'):
        label = 'Keywords:'
        content = text[len(label):]
    else:
        return False

    p = paragraph._element
    pPr = p.find(qn('w:pPr'))

    p.clear()
    if pPr is not None:
        p.append(pPr)

    r1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    b1 = OxmlElement('w:b')
    b1Cs = OxmlElement('w:bCs')
    rPr1.append(b1)
    rPr1.append(b1Cs)
    r1.append(rPr1)
    t1 = OxmlElement('w:t')
    t1.text = label
    r1.append(t1)
    p.append(r1)

    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = content
    r2.append(t2)
    p.append(r2)

    return True

def make_run_bold(paragraph):
    for run in paragraph.runs:
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        if rPr.find(qn('w:b')) is None:
            rPr.append(OxmlElement('w:b'))
        if rPr.find(qn('w:bCs')) is None:
            rPr.append(OxmlElement('w:bCs'))

def apply_style_by_id_para(element, style_id):
    pPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        element.insert(0, pPr)
    pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        if len(pPr) > 0:
            pPr.insert(0, pStyle)
        else:
            pPr.append(pStyle)
    pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', style_id)

def get_element_info(element):
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    if tag != 'p':
        return None, None, None
    pPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    raw_style_id = None
    if pPr is not None:
        pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        if pStyle is not None:
            raw_style_id = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
    text = ''
    for t in element.iter():
        if t.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            text += t.text or ''
    text = text.strip()
    current_style_name = None
    for para in doc.paragraphs:
        if para._element is element:
            current_style_name = para.style.name if para.style else 'Normal'
            break
    return raw_style_id, current_style_name, text

template_toc_styles = get_template_toc_style_ids(str(TEMPLATE_PATH))
if template_toc_styles:
    print(f"从模板读取目录样式: {template_toc_styles}")
    if '目录 一级标题' in template_toc_styles:
        TOC_STYLE_ID_MAPPING['14'] = template_toc_styles['目录 一级标题']
    if '目录 二级标题' in template_toc_styles:
        TOC_STYLE_ID_MAPPING['17'] = template_toc_styles['目录 二级标题']
    if '目录 三级标题' in template_toc_styles:
        TOC_STYLE_ID_MAPPING['8'] = template_toc_styles['目录 三级标题']

headings = []
for para in doc.paragraphs:
    text = para.text.strip()
    is_chapter = False
    level = 1
    
    if text.startswith('第') and '章' in text:
        is_chapter = True
        level = 1
    elif re.match(r'^\d+\.\d+\.\d+', text):
        is_chapter = True
        level = 3
    elif re.match(r'^\d+\.\d+', text):
        is_chapter = True
        level = 2
    
    if is_chapter:
        headings.append({
            'text': text,
            'level': level
        })

body = doc._element.body
sdt_element = None
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'sdt':
        sdt_element = child
        break

toc_generated = False
if sdt_element is not None and headings:
    sdtContent = sdt_element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent')
    if sdtContent is not None:
        toc_para = None
        for child in sdtContent:
            child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if child_tag == 'p':
                toc_para = child
        
        if toc_para is not None:
            for heading in headings:
                level = heading['level']
                text = heading['text']
                
                if level == 1:
                    style_id = template_toc_styles.get('目录 一级标题', '68')
                elif level == 2:
                    style_id = template_toc_styles.get('目录 二级标题', '69')
                else:
                    style_id = template_toc_styles.get('目录 三级标题', '70')
                
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '8306')
                tab.set(qn('w:leader'), 'dot')
                tabs.append(tab)
                pPr.append(tabs)
                
                pStyle = OxmlElement('w:pStyle')
                pStyle.set(qn('w:val'), style_id)
                pPr.append(pStyle)
                p.append(pPr)
                
                hl = OxmlElement('w:hyperlink')
                anchor = '_Toc' + str(hash(text) % 10000)
                hl.set(qn('w:anchor'), anchor)
                
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = text
                r.append(t)
                hl.append(r)
                p.append(hl)
                
                r_tab = OxmlElement('w:r')
                tab_char = OxmlElement('w:tab')
                r_tab.append(tab_char)
                p.append(r_tab)
                
                r_page = OxmlElement('w:r')
                t_page = OxmlElement('w:t')
                t_page.text = '1'
                r_page.append(t_page)
                p.append(r_page)
                
                toc_para.addprevious(p)
            
            toc_generated = True
            print(f"已生成 {len(headings)} 个目录条目")
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'sdt':
        for sdt_child in child:
            sdt_tag = sdt_child.tag.split('}')[-1] if '}' in sdt_child.tag else sdt_child.tag
            if sdt_tag == 'sdtContent':
                for content_child in sdt_child:
                    content_tag = content_child.tag.split('}')[-1] if '}' in content_child.tag else content_child.tag
                    if content_tag == 'p':
                        raw_style_id, current_style_name, text = get_element_info(content_child)
                        if raw_style_id in TOC_STYLE_ID_MAPPING:
                            target_style_id = TOC_STYLE_ID_MAPPING[raw_style_id]
                            apply_style_by_id_para(content_child, target_style_id)
                            print(f"目录ID映射: raw='{raw_style_id}' -> styleId={target_style_id}")
                        elif raw_style_id in TOC_STYLE_MAPPING:
                            target_style_id = TOC_STYLE_MAPPING[raw_style_id]
                            apply_style_by_id_para(content_child, target_style_id)
                            if text == '目录':
                                for r in content_child.iter():
                                    if r.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                                        rPr = r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                                        if rPr is None:
                                            rPr = OxmlElement('w:rPr')
                                            r.insert(0, rPr)
                                        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                                        if rFonts is None:
                                            rFonts = OxmlElement('w:rFonts')
                                            rPr.append(rFonts)
                                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Times New Roman')
                                        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Times New Roman')
                                        break
                                for t in content_child.iter():
                                    if t.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                                        t.text = '目    录'
                                        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                                        break
                            print(f"目录映射: raw='{raw_style_id}' -> styleId={target_style_id}")
    elif tag == 'p':
        raw_style_id, current_style_name, text = get_element_info(child)
        if raw_style_id is None:
            continue
        para_idx = None
        for idx, para in enumerate(doc.paragraphs):
            if para._element is child:
                para_idx = idx
                break
        abstract_context = None
        if para_idx is not None:
            abstract_context = detect_abstract_context(doc.paragraphs, para_idx)
        if abstract_context:
            context_type, target_style_id = abstract_context
            apply_style_by_id_para(child, target_style_id)
            if context_type == 'keywords':
                for para in doc.paragraphs:
                    if para._element is child:
                        split_keyword_paragraph(para)
                        break
            elif context_type == 'title' and text == 'Abstract':
                for para in doc.paragraphs:
                    if para._element is child:
                        make_run_bold(para)
                        break
            print(f"摘要映射: {context_type}='{text[:20]}' -> styleId={target_style_id}")
        elif raw_style_id in STYLE_MAPPING:
            target_style_id = STYLE_MAPPING[raw_style_id]
            apply_style_by_id_para(child, target_style_id)
            print(f"映射: raw='{raw_style_id}' -> styleId={target_style_id}")
        elif current_style_name and current_style_name in STYLE_MAPPING:
            target_style_id = STYLE_MAPPING[current_style_name]
            apply_style_by_id_para(child, target_style_id)
            print(f"映射: name='{current_style_name}' -> styleId={target_style_id}")
        elif raw_style_id in TOC_STYLE_ID_MAPPING:
            target_style_id = TOC_STYLE_ID_MAPPING[raw_style_id]
            apply_style_by_id_para(child, target_style_id)
            print(f"目录ID映射: raw='{raw_style_id}' -> styleId={target_style_id}")
        elif current_style_name and current_style_name in TOC_LEVEL_STYLE_MAPPING:
            target_style_id = TOC_LEVEL_STYLE_MAPPING[current_style_name]
            apply_style_by_id_para(child, target_style_id)
            print(f"目录级别映射: name='{current_style_name}' -> styleId={target_style_id}")
        else:
            if current_style_name and not current_style_name.startswith('Normal'):
                template_style_id = get_style_id_by_name(doc, current_style_name)
                if template_style_id:
                    apply_style_by_id_para(child, template_style_id)
                    print(f"保留: '{current_style_name}' -> styleId={template_style_id}")

for para in doc.paragraphs:
    p = para._element
    pPr = p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        if numPr is not None:
            ilvl = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            if ilvl is not None:
                ilvl_val = int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                
                ind = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind')
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                
                left_val = str(420 * (ilvl_val + 1))
                ind.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left', left_val)
                ind.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging', '420')
                
                print(f"列表缩进修复: ilvl={ilvl_val} left={left_val} hanging=420")

ref_section_started = False
ack_section_started = False
for para in doc.paragraphs:
    text = para.text.strip()
    p = para._element
    pPr = p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    
    if text == '参考文献':
        apply_style_by_id_para(p, '78')
        print(f"参考文献标题映射: '参考文献' -> styleId=78")
        ref_section_started = True
        ack_section_started = False
    elif text == '致谢':
        apply_style_by_id_para(p, '80')
        for t in p.iter():
            if t.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                t.text = '致    谢'
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                break
        print(f"致谢标题映射: '致谢' -> '致    谢' styleId=80")
        ack_section_started = True
        ref_section_started = False
    elif ref_section_started and text.startswith('['):
        if pPr is not None:
            pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pStyle is not None:
                current_style = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if current_style != '79':
                    apply_style_by_id_para(p, '79')
                    print(f"参考文献内容映射: '{text[:30]}' -> styleId=79")
    elif ack_section_started and text and not text.startswith('第') and '章' not in text:
        if pPr is not None:
            pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pStyle is not None:
                current_style = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if current_style != '81':
                    apply_style_by_id_para(p, '81')
                    print(f"致谢内容映射: '{text[:30]}' -> styleId=81")

doc.save(filename)
print(f"样式修复完成: {filename}")

import sys
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_size):
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(font_size))
    
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(font_size))

def fix_paragraph_format(paragraph, font_size, is_bold=False):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    
    pJc = pPr.find(qn('w:jc'))
    if pJc is None:
        pJc = OxmlElement('w:jc')
        pPr.append(pJc)
    pJc.set(qn('w:val'), 'center')
    
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    
    for run in paragraph.runs:
        set_run_font(run, font_size)
        if is_bold:
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr.find(qn('w:b')) is None:
                rPr.append(OxmlElement('w:b'))
            if rPr.find(qn('w:bCs')) is None:
                rPr.append(OxmlElement('w:bCs'))

def is_table_caption(text):
    return bool(re.match(r'^表\s*\d+', text.strip())) or bool(re.match(r'^Table\s*\d+', text.strip(), re.IGNORECASE))

def fix_tables(filename):
    doc = Document(filename)
    body = doc._element.body
    children = list(body)
    
    for i, child in enumerate(children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            idx = i - 1
            while idx >= 0:
                sibling = children[idx]
                sibling_tag = sibling.tag.split('}')[-1] if '}' in sibling.tag else sibling.tag
                if sibling_tag == 'p':
                    text = ''
                    for t in sibling.iter():
                        if t.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                            text += t.text or ''
                    text = text.strip()
                    if text and len(text) < 200:
                        for para in doc.paragraphs:
                            if para._element is sibling:
                                fix_paragraph_format(para, 24)
                                break
                    break
                idx -= 1
    
    for table in doc.tables:
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblStyle = tblPr.find(qn('w:tblStyle'))
        if tblStyle is not None:
            tblPr.remove(tblStyle)
        
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), '0')
        tblW.set(qn('w:type'), 'auto')
        
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')
        
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        else:
            for child in list(tblBorders):
                tblBorders.remove(child)
        
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:color'), 'auto')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:space'), '0')
        tblBorders.append(top)
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:color'), 'auto')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '0')
        tblBorders.append(bottom)
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'none')
        left.set(qn('w:color'), 'auto')
        left.set(qn('w:sz'), '0')
        left.set(qn('w:space'), '0')
        tblBorders.append(left)
        
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'none')
        right.set(qn('w:color'), 'auto')
        right.set(qn('w:sz'), '0')
        right.set(qn('w:space'), '0')
        tblBorders.append(right)
        
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'single')
        insideH.set(qn('w:color'), 'auto')
        insideH.set(qn('w:sz'), '4')
        insideH.set(qn('w:space'), '0')
        tblBorders.append(insideH)
        
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'none')
        insideV.set(qn('w:color'), 'auto')
        insideV.set(qn('w:sz'), '0')
        insideV.set(qn('w:space'), '0')
        tblBorders.append(insideV)
        
        tblCellMar = tblPr.find(qn('w:tblCellMar'))
        if tblCellMar is None:
            tblCellMar = OxmlElement('w:tblCellMar')
            tblPr.append(tblCellMar)
        else:
            for child in list(tblCellMar):
                tblCellMar.remove(child)
        
        for side in ['top', 'left', 'bottom', 'right']:
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'), '108')
            mar.set(qn('w:type'), 'dxa')
            tblCellMar.append(mar)
        
        for row_idx, row in enumerate(table.rows):
            tr = row._element
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            
            trJc = trPr.find(qn('w:jc'))
            if trJc is None:
                trJc = OxmlElement('w:jc')
                trPr.append(trJc)
            trJc.set(qn('w:val'), 'center')
            
            tblPrEx = tr.find(qn('w:tblPrEx'))
            if tblPrEx is None:
                tblPrEx = OxmlElement('w:tblPrEx')
                tr.insert(0, tblPrEx)
            
            exBorders = tblPrEx.find(qn('w:tblBorders'))
            if exBorders is None:
                exBorders = OxmlElement('w:tblBorders')
                tblPrEx.append(exBorders)
            else:
                for child in list(exBorders):
                    exBorders.remove(child)
            
            for border_name, val, sz in [
                ('top', 'single', '12'),
                ('bottom', 'single', '12'),
                ('left', 'none', '0'),
                ('right', 'none', '0'),
                ('insideH', 'single', '4'),
                ('insideV', 'none', '0')
            ]:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), val)
                border.set(qn('w:color'), 'auto')
                border.set(qn('w:sz'), sz)
                border.set(qn('w:space'), '0')
                exBorders.append(border)
            
            exCellMar = tblPrEx.find(qn('w:tblCellMar'))
            if exCellMar is None:
                exCellMar = OxmlElement('w:tblCellMar')
                tblPrEx.append(exCellMar)
            else:
                for child in list(exCellMar):
                    exCellMar.remove(child)
            
            for side, w in [('top', '0'), ('left', '108'), ('bottom', '0'), ('right', '108')]:
                mar = OxmlElement(f'w:{side}')
                mar.set(qn('w:w'), w)
                mar.set(qn('w:type'), 'dxa')
                exCellMar.append(mar)
            
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                else:
                    for child in list(tcBorders):
                        tcBorders.remove(child)
                
                is_first_row = (row_idx == 0)
                is_last_row = (row_idx == len(table.rows) - 1)
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    
                    if border_name == 'top':
                        if is_first_row:
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:color'), 'auto')
                            border.set(qn('w:sz'), '12')
                            border.set(qn('w:space'), '0')
                        else:
                            border.set(qn('w:val'), 'nil')
                    elif border_name == 'bottom':
                        if is_last_row:
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:color'), 'auto')
                            border.set(qn('w:sz'), '12')
                            border.set(qn('w:space'), '0')
                        else:
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:color'), 'auto')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:space'), '0')
                    else:
                        border.set(qn('w:val'), 'nil')
                    
                    tcBorders.append(border)
                
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = OxmlElement('w:vAlign')
                    tcPr.append(vAlign)
                vAlign.set(qn('w:val'), 'center')
                
                for paragraph in cell.paragraphs:
                    pPr = paragraph._element.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        paragraph._element.insert(0, pPr)
                    
                    pJc = pPr.find(qn('w:jc'))
                    if pJc is None:
                        pJc = OxmlElement('w:jc')
                        pPr.append(pJc)
                    pJc.set(qn('w:val'), 'center')
                    
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:line'), '360')
                    spacing.set(qn('w:lineRule'), 'auto')
                    
                    for run in paragraph.runs:
                        r = run._element
                        rPr = r.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.insert(0, rPr)
                        
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        rFonts.set(qn('w:eastAsia'), '宋体')
                        rFonts.set(qn('w:cs'), 'Times New Roman')
                        
                        sz = rPr.find(qn('w:sz'))
                        if sz is None:
                            sz = OxmlElement('w:sz')
                            rPr.append(sz)
                        sz.set(qn('w:val'), '21')
                        
                        szCs = rPr.find(qn('w:szCs'))
                        if szCs is None:
                            szCs = OxmlElement('w:szCs')
                            rPr.append(szCs)
                        szCs.set(qn('w:val'), '21')
                        
                        b = rPr.find(qn('w:b'))
                        if b is not None:
                            rPr.remove(b)
                        bCs = rPr.find(qn('w:bCs'))
                        if bCs is not None:
                            rPr.remove(bCs)
    
    doc.save(filename)
    print(f"表格格式修复完成: {filename}")

if __name__ == '__main__':
    fix_tables(sys.argv[1])

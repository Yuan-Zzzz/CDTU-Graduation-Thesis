from pathlib import Path

from docx import Document
from docxcompose.composer import Composer

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
OUTPUT_DIR = REPO_ROOT / "output"

master = Document(str(OUTPUT_DIR / "abstract.docx"))
composer = Composer(master)

doc2 = Document(str(OUTPUT_DIR / "thesis.docx"))
composer.append(doc2)

doc3 = Document(str(OUTPUT_DIR / "acknowledgments.docx"))
composer.append(doc3)

master.save(str(OUTPUT_DIR / "Final_thesis.docx"))
